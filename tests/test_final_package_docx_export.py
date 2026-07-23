from __future__ import annotations

from io import BytesIO

from docx import Document

from exporters.final_package_exporter import FinalPackageExporter
from tests.final_package_helpers import build_package, fake_banner_bytes


def test_docx_export_opens_contains_sections_tables_footer_and_unicode():
    data = FinalPackageExporter().export_docx(build_package(banner=True), banner_image_bytes=fake_banner_bytes())
    document = Document(BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert data.startswith(b"PK")
    assert "1. Portada" in text
    assert "16. Metodología y disclaimer" in text
    assert "Gestión" in text
    assert document.tables
    assert document.sections[0].footer.paragraphs[0].text


def test_docx_export_works_without_banner():
    data = FinalPackageExporter().export_docx(build_package(banner=False))
    document = Document(BytesIO(data))

    assert document.paragraphs
    assert len(data) > 1000
