from __future__ import annotations

from io import BytesIO

from docx import Document

from exporters.targeted_cv_docx_exporter import TargetedCVDocxExporter
from services.export_audit_service import ExportAuditService
from tests.targeted_cv_helpers import build_targeted_cv


def test_targeted_cv_docx_export_opens_without_tables_or_images():
    data = TargetedCVDocxExporter().export(build_targeted_cv(1))
    document = Document(BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "PERFIL PROFESIONAL" in text
    assert "EXPERIENCIA PROFESIONAL" in text
    assert document.tables == []
    assert len(document.inline_shapes) == 0
    assert ExportAuditService().audit_targeted_cv_docx(data).passed
