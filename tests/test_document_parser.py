from __future__ import annotations

from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from services.candidate_input_service import prepare_candidate_input
from services.document_parser import (
    DOCX_ERROR,
    PDF_OPEN_ERROR,
    PDF_SCANNED_ERROR,
    PDF_TEXT_ERROR,
    extract_docx_text,
    extract_pdf_text,
    normalize_extracted_text,
    parse_uploaded_document,
)
from utils.validators import JobFormInput


LONG_CV_LINE = (
    "María Ejemplo es Project Manager con experiencia ficticia en gestión de proyectos, "
    "coordinación de equipos, seguimiento de indicadores, comunicación ejecutiva y mejora "
    "continua de procesos para organizaciones tecnológicas. "
)

VALID_JOB_DESCRIPTION = (
    "Responsable de liderar iniciativas estratégicas, coordinar equipos multifuncionales, "
    "analizar métricas de negocio y comunicar prioridades con stakeholders internos."
)


def test_normalize_extracted_text_preserves_useful_text():
    raw_text = "  María  Ejemplo\r\n\x00• Liderazgo   técnico\r\r\n\n\n\nExperiencia limpia  "

    assert normalize_extracted_text(raw_text) == (
        "María Ejemplo\n"
        "• Liderazgo técnico\n\n"
        "Experiencia limpia"
    )


def test_normalize_extracted_text_keeps_clean_text():
    text = "Perfil profesional\n\nExperiencia en producto"

    assert normalize_extracted_text(text) == text


def test_docx_extracts_paragraphs():
    result = extract_docx_text(_make_docx(paragraphs=[LONG_CV_LINE, LONG_CV_LINE]), "cv.docx")

    assert result.success
    assert result.file_type == "docx"
    assert result.paragraph_count == 2
    assert "María Ejemplo" in result.normalized_text
    assert result.word_count > 20


def test_docx_extracts_table_rows():
    result = extract_docx_text(
        _make_docx(
            table=[
                ["Empresa", "Cargo", "Periodo"],
                ["Acme", "Project Manager", "2020-2026"],
            ]
        ),
        "tabla.docx",
    )

    assert "Empresa | Cargo | Periodo" in result.normalized_text
    assert "Acme | Project Manager | 2020-2026" in result.normalized_text


def test_docx_extracts_paragraphs_and_table():
    result = extract_docx_text(
        _make_docx(paragraphs=[LONG_CV_LINE, LONG_CV_LINE], table=[["Skill", "Nivel"], ["PM", "Avanzado"]]),
        "mixto.docx",
    )

    assert result.success
    assert "Skill | Nivel" in result.normalized_text
    assert result.paragraph_count == 2


def test_docx_empty_file_is_error():
    result = extract_docx_text(_make_docx(), "vacio.docx")

    assert not result.success
    assert DOCX_ERROR in result.errors


def test_docx_with_little_text_warns():
    result = extract_docx_text(_make_docx(paragraphs=["Texto breve."]), "breve.docx")

    assert result.success
    assert result.warnings
    assert result.character_count < 200


def test_docx_invalid_bytes_are_error():
    result = extract_docx_text(b"no es un docx", "corrupto.docx")

    assert not result.success
    assert DOCX_ERROR in result.errors


def test_pdf_one_page_with_text():
    result = extract_pdf_text(_make_pdf([LONG_CV_LINE * 2]), "cv.pdf")

    assert result.success
    assert result.file_type == "pdf"
    assert result.page_count == 1
    assert "María Ejemplo" in result.normalized_text
    assert result.character_count >= 200


def test_pdf_multiple_pages_with_text():
    result = extract_pdf_text(_make_pdf([LONG_CV_LINE * 2, LONG_CV_LINE * 2]), "multi.pdf")

    assert result.success
    assert result.page_count == 2
    assert result.word_count > 40


def test_pdf_with_empty_page_warns_but_can_succeed():
    result = extract_pdf_text(_make_pdf([LONG_CV_LINE * 3, ""]), "parcial.pdf")

    assert result.success
    assert result.page_count == 2
    assert any("no devolvieron texto" in warning for warning in result.warnings)


def test_pdf_without_enough_text_is_error():
    result = extract_pdf_text(_make_pdf(["Texto digital pero insuficiente para análisis. " * 2]), "breve.pdf")

    assert not result.success
    assert PDF_TEXT_ERROR in result.errors
    assert not result.likely_scanned


def test_pdf_likely_scanned_is_error():
    result = extract_pdf_text(_make_pdf(["", "", ""]), "escaneado.pdf")

    assert not result.success
    assert result.likely_scanned
    assert PDF_SCANNED_ERROR in result.errors


def test_pdf_invalid_bytes_are_error():
    result = extract_pdf_text(b"no es pdf", "corrupto.pdf")

    assert not result.success
    assert PDF_OPEN_ERROR in result.errors


def test_parse_uploaded_document_detects_type_by_extension():
    result = parse_uploaded_document("cv.docx", _make_docx(paragraphs=[LONG_CV_LINE, LONG_CV_LINE]))

    assert result.file_type == "docx"
    assert result.success


def test_prepare_candidate_input_with_docx_file_prepares_cv():
    docx_bytes = _make_docx(paragraphs=[LONG_CV_LINE, LONG_CV_LINE])
    result = prepare_candidate_input(
        cv_text="",
        cv_file_name="cv.docx",
        cv_file_size=len(docx_bytes),
        cv_file_bytes=docx_bytes,
        linkedin_text="",
        linkedin_url="",
        jobs=[
            JobFormInput(index=1, title="Product Manager", description=VALID_JOB_DESCRIPTION),
            JobFormInput(index=2, title="Program Manager", description=VALID_JOB_DESCRIPTION),
        ],
        output_language="es",
    )

    assert result.is_valid
    assert result.cv_source == "docx"
    assert result.cv_text.startswith("María Ejemplo")
    assert result.cv_summary.file_type == "docx"


def test_prepare_candidate_input_with_text_does_not_parse_file():
    result = prepare_candidate_input(
        cv_text=LONG_CV_LINE * 2,
        cv_file_name="cv.pdf",
        cv_file_size=8,
        cv_file_bytes=b"bad pdf bytes",
        linkedin_text="",
        linkedin_url="",
        jobs=[
            JobFormInput(index=1, title="Product Manager", description=VALID_JOB_DESCRIPTION),
            JobFormInput(index=2, title="Program Manager", description=VALID_JOB_DESCRIPTION),
        ],
        output_language="en",
    )

    assert result.is_valid
    assert result.cv_source == "text"
    assert result.cv_parse_result is None
    assert any("no fue necesario" in message.message for message in result.messages)


def test_prepare_candidate_input_with_short_text_is_invalid():
    result = prepare_candidate_input(
        cv_text="CV corto.",
        cv_file_name=None,
        cv_file_size=None,
        cv_file_bytes=None,
        linkedin_text="",
        linkedin_url="",
        jobs=[
            JobFormInput(index=1, title="Product Manager", description=VALID_JOB_DESCRIPTION),
            JobFormInput(index=2, title="Program Manager", description=VALID_JOB_DESCRIPTION),
        ],
        output_language="es",
    )

    assert not result.is_valid
    assert result.cv_text is None


def _make_docx(
    paragraphs: list[str] | None = None,
    table: list[list[str]] | None = None,
) -> bytes:
    document = Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table:
        doc_table = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for col_index, value in enumerate(row):
                doc_table.cell(row_index, col_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf(pages: list[str]) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    for page_text in pages:
        text_object = pdf.beginText(72, 740)
        for line in _wrap_text(page_text):
            text_object.textLine(line)
        pdf.drawText(text_object)
        pdf.showPage()

    pdf.save()
    return buffer.getvalue()


def _wrap_text(text: str) -> list[str]:
    if not text:
        return []

    words = text.split()
    lines = []
    current = []
    for word in words:
        current.append(word)
        if len(" ".join(current)) >= 80:
            lines.append(" ".join(current))
            current = []

    if current:
        lines.append(" ".join(current))
    return lines
