"""DOCX exporter for application communication kits."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

from schemas.application_communication_models import ApplicationCommunicationKit


class ApplicationCommunicationDocxExporter:
    """Export one communication kit to a simple DOCX document."""

    def export(self, kit: ApplicationCommunicationKit) -> bytes:
        document = Document()
        _setup_document(document)
        _add_title(document, kit)
        _add_heading(document, "Carta de presentación")
        _add_paragraphs(document, kit.cover_letter.full_text)
        _add_heading(document, "Mensaje para recruiter")
        _add_paragraphs(document, kit.recruiter_message.message)
        _add_heading(document, "Asuntos sugeridos")
        for subject in kit.application_email.subject_options:
            _add_bullet(document, subject)
        _add_heading(document, "Correo de postulación")
        _add_paragraphs(document, kit.application_email.full_text)
        _add_heading(document, "Notas de revisión")
        for note in _review_notes(kit):
            _add_bullet(document, note)
        _add_disclaimer(document)
        _scrub_properties(document)
        output = BytesIO()
        document.save(output)
        return output.getvalue()


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)


def _add_title(document: Document, kit: ApplicationCommunicationKit) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Kit de postulación - Vacante {kit.target_job_index:02d}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    details = document.add_paragraph()
    details.add_run(f"{kit.target_job_title} | {kit.target_company or 'Empresa no especificada'}")


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def _add_paragraphs(document: Document, text: str) -> None:
    for block in [value.strip() for value in text.splitlines() if value.strip()]:
        paragraph = document.add_paragraph(block)
        paragraph.paragraph_format.space_after = Pt(4)


def _add_bullet(document: Document, text: str) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def _add_disclaimer(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    run = paragraph.add_run("Contenido orientativo; revisar antes de enviar.")
    run.italic = True
    run.font.size = Pt(9)


def _review_notes(kit: ApplicationCommunicationKit) -> list[str]:
    notes = [
        *kit.personalization_notes,
        *kit.cover_letter.review_notes,
        *kit.recruiter_message.review_notes,
        *kit.application_email.review_notes,
        *kit.risks_or_claims_requiring_review,
    ]
    return [note for note in notes if str(note).strip()] or ["Revisar datos, tono y adjuntos antes de enviar."]


def _scrub_properties(document: Document) -> None:
    props = document.core_properties
    props.author = ""
    props.comments = ""
    props.keywords = ""
    props.subject = ""
    props.title = ""
