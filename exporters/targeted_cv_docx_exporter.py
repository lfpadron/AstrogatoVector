"""DOCX exporter for targeted CVs."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from schemas.targeted_cv_models import TargetedCV


class TargetedCVDocxExporter:
    """Export one targeted CV to a simple one-column DOCX."""

    def export(self, targeted_cv: TargetedCV) -> bytes:
        document = Document()
        _setup_document(document)
        _add_header(document, targeted_cv)
        _add_heading(document, "Perfil profesional")
        _add_paragraph(document, targeted_cv.summary.text)
        _add_heading(document, "Competencias clave")
        for skill in sorted(targeted_cv.skills, key=lambda item: item.priority):
            _add_bullet(document, skill.name)
        _add_heading(document, "Experiencia profesional")
        for entry in targeted_cv.experience:
            if not entry.included:
                continue
            _add_subheading(document, f"{entry.display_role_title} | {entry.employer}")
            details = " | ".join(value for value in [_date_line(entry.start_date, entry.end_date), entry.location] if value)
            if details:
                _add_small(document, details)
            for bullet in entry.bullets:
                _add_bullet(document, bullet.text)
        _add_optional_section(document, "Educación", [entry.text for entry in targeted_cv.education if entry.visible])
        _add_optional_section(
            document,
            "Certificaciones",
            [entry.text for entry in targeted_cv.certifications if entry.visible],
        )
        _add_optional_section(document, "Idiomas", [entry.text for entry in targeted_cv.languages if entry.visible])
        _scrub_properties(document)
        output = BytesIO()
        document.save(output)
        return output.getvalue()


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)


def _add_header(document: Document, targeted_cv: TargetedCV) -> None:
    header = targeted_cv.header
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(header.candidate_name or header.professional_title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(header.professional_title)
    contact = " | ".join(value for value in [header.location, header.email, header.phone, header.linkedin_url] if value)
    if contact:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(contact)


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def _add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_small(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _add_bullet(document: Document, text: str) -> None:
    if not text.strip():
        return
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run(text)


def _add_optional_section(document: Document, title: str, values: list[str]) -> None:
    if not values:
        return
    _add_heading(document, title)
    for value in values:
        _add_bullet(document, value)


def _date_line(start_date: str | None, end_date: str | None) -> str:
    if start_date and end_date:
        return f"{start_date} - {end_date}"
    if start_date:
        return f"{start_date} - Actual"
    return end_date or ""


def _scrub_properties(document: Document) -> None:
    props = document.core_properties
    props.author = ""
    props.comments = ""
    props.keywords = ""
    props.subject = ""
    props.title = ""
