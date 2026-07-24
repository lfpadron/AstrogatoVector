"""DOCX exporter for LinkedIn professional editorial plans."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

from schemas.editorial_plan_models import LinkedInPostPlan, ProfessionalBrandPlan
from utils.filename_utils import safe_download_filename

EDITORIAL_PLAN_DOCX_FILENAME = safe_download_filename("linkedin-editorial-plan", "docx")


class EditorialPlanDocxExporter:
    """Export one professional brand plan to DOCX bytes."""

    def export(self, plan: ProfessionalBrandPlan) -> bytes:
        """Return DOCX bytes."""
        document = Document()
        _setup_document(document)
        _add_title(document, "Plan editorial profesional de LinkedIn")
        _add_paragraph(document, "Borrador profesional para revisión humana y publicación manual.")
        _add_heading(document, "Resumen", level=1)
        _add_paragraph(document, plan.summary)
        _add_list_section(document, "Objetivos", [str(item) for item in plan.objectives])
        _add_list_section(document, "Fortalezas explotadas", plan.strengths_exploited)
        _add_list_section(document, "Temas", plan.themes)
        _add_list_section(document, "Riesgos", plan.risks)
        _add_list_section(document, "Recomendaciones", plan.recommendations)
        _add_heading(document, "Calendario editorial", level=1)
        for week in sorted(plan.calendar.weeks, key=lambda item: item.week):
            _add_heading(document, f"Semana {week.week}", level=2)
            for post in sorted(week.posts, key=_day_order):
                _add_post(document, post)
        _add_heading(document, "Nota de uso", level=1)
        _add_paragraph(document, "Estos textos son orientativos. Revisa evidencia, tono y confidencialidad antes de publicarlos.")
        _scrub_properties(document)
        output = BytesIO()
        document.save(output)
        return output.getvalue()


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)


def _add_heading(document: Document, text: str, *, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14 if level == 1 else 12)


def _add_post(document: Document, post: LinkedInPostPlan) -> None:
    _add_heading(document, f"{_day_label(post.day)} - {post.title}", level=2)
    _add_paragraph(
        document,
        (
            f"Semana {post.week} | Objetivo: {post.objective} | Tipo: {post.post_type} | "
            f"Formato: {post.format} | Caracteres: {post.character_count}"
        ),
    )
    _add_paragraph(document, f"Tema: {post.theme}")
    _add_paragraph(document, f"Audiencia: {post.audience}")
    _add_label(document, "Hook")
    _add_paragraph(document, post.hook)
    _add_label(document, "Texto")
    _add_paragraphs(document, post.body)
    _add_label(document, "CTA")
    _add_paragraph(document, post.cta)
    _add_label(document, "Hashtags")
    _add_paragraph(document, " ".join(post.hashtags) if post.hashtags else "Sin hashtags")
    _add_list_section(document, "Keywords utilizadas", post.keywords_used, level=3)
    _add_list_section(document, "Evidencia utilizada", post.evidence_used, level=3)
    _add_list_section(document, "Claims que requieren revisión", post.claims_requiring_review, level=3)
    _add_list_section(document, "Notas", post.notes, level=3)


def _add_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True


def _add_paragraphs(document: Document, text: str) -> None:
    for block in [value.strip() for value in text.splitlines() if value.strip()]:
        _add_paragraph(document, block)


def _add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(3)


def _add_list_section(document: Document, title: str, values: list[object], *, level: int = 2) -> None:
    _add_heading(document, title, level=1 if level <= 2 else 2)
    items = [str(value).strip() for value in values if str(value).strip()] or ["Sin elementos."]
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def _scrub_properties(document: Document) -> None:
    props = document.core_properties
    props.author = ""
    props.comments = ""
    props.keywords = ""
    props.subject = ""
    props.title = ""


def _day_order(post: LinkedInPostPlan) -> int:
    return {"monday": 0, "wednesday": 1, "friday": 2}.get(str(post.day), 99)


def _day_label(value: object) -> str:
    labels = {
        "monday": "Lunes",
        "wednesday": "Miércoles",
        "friday": "Viernes",
    }
    return labels.get(str(value), str(value))
