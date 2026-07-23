"""In-memory exporters for the final professional package."""

from __future__ import annotations

import base64
import html
import json
import zipfile
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from schemas.audit_models import FINAL_AUDIT_DISCLAIMER
from schemas.compatibility_models import COMPATIBILITY_BAND_LABELS_EN, COMPATIBILITY_BAND_LABELS_ES
from schemas.deliverable_models import (
    FINAL_PACKAGE_DISCLAIMERS,
    FINAL_PACKAGE_EXPORT_VERSION,
    FINAL_PACKAGE_REQUIRED_SECTIONS,
    FinalDeliverablePackage,
)
from schemas.enums import OutputLanguage
from services.font_service import DEFAULT_FONT_NAMES, find_font_path
from utils.filename_utils import safe_download_filename

INDIVIDUAL_MARKDOWN_FILENAME = safe_download_filename("astrogato-vector-perfil-profesional", "md")
INDIVIDUAL_HTML_FILENAME = safe_download_filename("astrogato-vector-perfil-profesional", "html")
INDIVIDUAL_DOCX_FILENAME = safe_download_filename("astrogato-vector-perfil-profesional", "docx")
INDIVIDUAL_PDF_FILENAME = safe_download_filename("astrogato-vector-perfil-profesional", "pdf")
FINAL_ZIP_FILENAME = safe_download_filename("astrogato-vector-paquete-profesional", "zip")
BANNER_ZIP_FILENAME = "linkedin-banner.png"
ZIP_ROOT = "astrogato-vector-paquete-profesional"


class FinalPackageExporter:
    """Export a package to Markdown, HTML, DOCX, PDF and ZIP bytes."""

    def export_markdown(self, package: FinalDeliverablePackage, *, banner_filename: str = BANNER_ZIP_FILENAME) -> bytes:
        """Export Markdown as UTF-8 bytes."""
        banner_line = f"![Banner de LinkedIn]({banner_filename})" if package.banner_included else ""
        lines = [
            f"# {package.package_title}",
            "",
            _metadata_line("Generated with", "Astrogato Vector"),
            _metadata_line(_label(package, "date"), _format_datetime(package.generated_at)),
            _metadata_line(_label(package, "language"), package.output_language),
            "",
            "## 1. Portada",
            "",
            f"**{_label(package, 'identity')}:** {package.professional_identity}",
            f"**{_label(package, 'roles')}:** {_join(package.target_roles)}",
            "",
            "## 2. Resumen ejecutivo",
            "",
            package.executive_summary,
            "",
            "## 3. Perfil de LinkedIn",
            "",
            "## 4. Banner textual",
            "",
            banner_line,
            "",
            f"- {_label(package, 'banner_primary')}: {package.banner_primary_line}",
            f"- {_label(package, 'banner_specialty')}: {package.banner_specialty_line}",
            f"- {_label(package, 'banner_supporting')}: {package.banner_supporting_line or _label(package, 'not_available')}",
            "",
            "## 5. Headline",
            "",
            package.headline,
            "",
            "## 6. About",
            "",
            package.about,
            "",
            "## 7. Experiencia profesional",
            "",
        ]
        for experience in package.experiences:
            lines.extend(
                [
                    f"### {experience.employer} - {experience.suggested_role_title}",
                    "",
                    experience.rewritten_text,
                    "",
                ]
            )
        lines.extend(
            [
                "## 8. Skills priorizadas",
                "",
                "| Rank | Skill | Categoría | Evidencia | Ubicación |",
                "| --- | --- | --- | --- | --- |",
            ]
        )
        for skill in sorted(package.prioritized_skills, key=lambda item: item.priority_rank):
            lines.append(
                "| "
                f"{skill.priority_rank} | {_md_cell(skill.name)} | {_md_cell(skill.category)} | "
                f"{_md_cell(skill.evidence_status)} | {_md_cell(_join(skill.recommended_placement))} |"
            )
        lines.extend(
            [
                "",
                "## 9. Keywords ATS",
                "",
                "| Keyword | Prioridad | Frecuencia | Respaldada | Secciones |",
                "| --- | --- | ---: | --- | --- |",
            ]
        )
        for keyword in package.ats_keywords:
            lines.append(
                "| "
                f"{_md_cell(keyword.keyword)} | {_md_cell(keyword.priority)} | {keyword.frequency_in_jobs} | "
                f"{_yes_no(package, keyword.supported_by_candidate)} | {_md_cell(_join(keyword.recommended_sections))} |"
            )
        lines.extend(_markdown_compatibility(package))
        lines.extend(_markdown_lists(package))
        lines.extend(_markdown_audits(package))
        lines.extend(_markdown_methodology(package))
        return "\n".join(line for line in lines if line is not None).encode("utf-8")

    def export_html(
        self,
        package: FinalDeliverablePackage,
        *,
        banner_image_bytes: bytes | None = None,
        embed_banner: bool = True,
        banner_filename: str = BANNER_ZIP_FILENAME,
    ) -> bytes:
        """Export self-contained or ZIP-relative HTML bytes."""
        banner_html = ""
        if package.banner_included and banner_image_bytes:
            src = (
                f"data:image/png;base64,{base64.b64encode(banner_image_bytes).decode('ascii')}"
                if embed_banner
                else banner_filename
            )
            banner_html = f'<img class="banner" src="{src}" alt="Banner de LinkedIn">'
        body = [
            '<!DOCTYPE html>',
            '<html lang="' + html.escape(str(package.output_language)) + '">',
            "<head>",
            '<meta charset="utf-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1">',
            "<title>" + html.escape(package.package_title) + "</title>",
            "<style>" + _html_css() + "</style>",
            "</head>",
            "<body>",
            "<main>",
            '<section class="cover">',
            "<p>Generated with Astrogato Vector</p>",
            "<h1>" + html.escape(package.package_title) + "</h1>",
            "<dl>",
            _html_kv(_label(package, "identity"), package.professional_identity),
            _html_kv(_label(package, "roles"), _join(package.target_roles)),
            _html_kv(_label(package, "date"), _format_datetime(package.generated_at)),
            "</dl>",
            "</section>",
            _html_section("1. Portada", ""),
            _html_section("2. Resumen ejecutivo", f"<p>{html.escape(package.executive_summary)}</p>"),
            _html_section("3. Perfil de LinkedIn", ""),
            _html_section(
                "4. Banner textual",
                banner_html
                + _html_list(
                    [
                        f"{_label(package, 'banner_primary')}: {package.banner_primary_line}",
                        f"{_label(package, 'banner_specialty')}: {package.banner_specialty_line}",
                        f"{_label(package, 'banner_supporting')}: {package.banner_supporting_line or _label(package, 'not_available')}",
                    ]
                ),
            ),
            _html_section("5. Headline", f"<p>{html.escape(package.headline)}</p>"),
            _html_section("6. About", f"<p>{html.escape(package.about)}</p>"),
            _html_section("7. Experiencia profesional", _html_experiences(package)),
            _html_section("8. Skills priorizadas", _html_skills_table(package)),
            _html_section("9. Keywords ATS", _html_keywords_table(package)),
            _html_section("10. Compatibilidad por vacante", _html_compatibility_table(package)),
            _html_section("11. Fortalezas", _html_list(package.key_strengths)),
            _html_section("12. Brechas", _html_list(package.critical_gaps)),
            _html_section("13. Recomendaciones", _html_list(package.strategic_recommendations)),
            _html_section("14. Auditoría LinkedIn", _html_audit_components(package, "linkedin")),
            _html_section("15. Auditoría ATS", _html_audit_components(package, "ats")),
            _html_section("16. Metodología y disclaimer", f"<p>{html.escape(_disclaimer(package))}</p>"),
            "</main>",
            "</body>",
            "</html>",
        ]
        return "\n".join(body).encode("utf-8")

    def export_docx(
        self,
        package: FinalDeliverablePackage,
        *,
        banner_image_bytes: bytes | None = None,
    ) -> bytes:
        """Export a DOCX report using python-docx."""
        document = Document()
        _setup_docx(document)
        section = document.sections[0]
        section.header.paragraphs[0].text = "Generated with Astrogato Vector"
        section.footer.paragraphs[0].text = _disclaimer(package)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(package.package_title)
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(20, 54, 92)
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(package.professional_identity)
        document.add_paragraph(_format_datetime(package.generated_at))
        document.add_page_break()

        _docx_heading(document, "1. Portada")
        _docx_paragraph(document, f"{_label(package, 'identity')}: {package.professional_identity}")
        _docx_paragraph(document, f"{_label(package, 'roles')}: {_join(package.target_roles)}")
        _docx_heading(document, "2. Resumen ejecutivo")
        _docx_paragraph(document, package.executive_summary)
        _docx_heading(document, "3. Perfil de LinkedIn")
        _docx_heading(document, "4. Banner textual", level=2)
        if package.banner_included and banner_image_bytes:
            _docx_add_banner(document, banner_image_bytes)
        _docx_bullets(
            document,
            [
                f"{_label(package, 'banner_primary')}: {package.banner_primary_line}",
                f"{_label(package, 'banner_specialty')}: {package.banner_specialty_line}",
                f"{_label(package, 'banner_supporting')}: {package.banner_supporting_line or _label(package, 'not_available')}",
            ],
        )
        _docx_heading(document, "5. Headline", level=2)
        _docx_paragraph(document, package.headline)
        _docx_heading(document, "6. About", level=2)
        _docx_paragraph(document, package.about)
        _docx_heading(document, "7. Experiencia profesional")
        for experience in package.experiences:
            _docx_heading(document, f"{experience.employer} - {experience.suggested_role_title}", level=2)
            _docx_paragraph(document, experience.rewritten_text)
        _docx_heading(document, "8. Skills priorizadas")
        _docx_table(
            document,
            ["Rank", "Skill", "Categoría", "Evidencia"],
            [
                [str(skill.priority_rank), skill.name, str(skill.category), str(skill.evidence_status)]
                for skill in sorted(package.prioritized_skills, key=lambda item: item.priority_rank)
            ],
        )
        _docx_heading(document, "9. Keywords ATS")
        _docx_table(
            document,
            ["Keyword", "Prioridad", "Frecuencia", "Respaldada"],
            [
                [
                    keyword.keyword,
                    str(keyword.priority),
                    str(keyword.frequency_in_jobs),
                    _yes_no(package, keyword.supported_by_candidate),
                ]
                for keyword in package.ats_keywords
            ],
        )
        _docx_heading(document, "10. Compatibilidad por vacante")
        _docx_table(document, ["Vacante", "Empresa", "Score", "Banda"], _compatibility_rows(package))
        _docx_heading(document, "11. Fortalezas")
        _docx_bullets(document, package.key_strengths)
        _docx_heading(document, "12. Brechas")
        _docx_bullets(document, package.critical_gaps)
        _docx_heading(document, "13. Recomendaciones")
        _docx_bullets(document, package.strategic_recommendations)
        _docx_heading(document, "14. Auditoría LinkedIn")
        _docx_audit(document, package, "linkedin")
        _docx_heading(document, "15. Auditoría ATS")
        _docx_audit(document, package, "ats")
        _docx_heading(document, "16. Metodología y disclaimer")
        _docx_paragraph(document, _disclaimer(package))
        _scrub_docx_core_properties(document)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def export_pdf(
        self,
        package: FinalDeliverablePackage,
        *,
        banner_image_bytes: bytes | None = None,
    ) -> bytes:
        """Export a selectable PDF using ReportLab."""
        output = BytesIO()
        doc = SimpleDocTemplate(
            output,
            pagesize=LETTER,
            rightMargin=0.72 * inch,
            leftMargin=0.72 * inch,
            topMargin=0.72 * inch,
            bottomMargin=0.72 * inch,
            title=package.package_title,
        )
        styles = _pdf_styles()
        story: list[Any] = [
            Paragraph(html.escape(package.package_title), styles["Title"]),
            Paragraph(html.escape(package.professional_identity), styles["Subtitle"]),
            Paragraph(f"Generated with Astrogato Vector - {_format_datetime(package.generated_at)}", styles["Small"]),
            Spacer(1, 0.25 * inch),
            Paragraph("1. Portada", styles["Heading1"]),
            Paragraph(html.escape(f"{_label(package, 'roles')}: {_join(package.target_roles)}"), styles["Body"]),
            Paragraph("2. Resumen ejecutivo", styles["Heading1"]),
            Paragraph(html.escape(package.executive_summary), styles["Body"]),
            Paragraph("3. Perfil de LinkedIn", styles["Heading1"]),
            Paragraph("4. Banner textual", styles["Heading2"]),
        ]
        if package.banner_included and banner_image_bytes:
            story.append(_pdf_banner(banner_image_bytes))
        story.extend(
            [
                Paragraph(html.escape(f"{_label(package, 'banner_primary')}: {package.banner_primary_line}"), styles["Body"]),
                Paragraph(html.escape(f"{_label(package, 'banner_specialty')}: {package.banner_specialty_line}"), styles["Body"]),
                Paragraph(
                    html.escape(
                        f"{_label(package, 'banner_supporting')}: {package.banner_supporting_line or _label(package, 'not_available')}"
                    ),
                    styles["Body"],
                ),
                Paragraph("5. Headline", styles["Heading2"]),
                Paragraph(html.escape(package.headline), styles["Body"]),
                Paragraph("6. About", styles["Heading2"]),
                Paragraph(html.escape(package.about), styles["Body"]),
                Paragraph("7. Experiencia profesional", styles["Heading1"]),
            ]
        )
        for experience in package.experiences:
            story.append(Paragraph(html.escape(f"{experience.employer} - {experience.suggested_role_title}"), styles["Heading2"]))
            story.append(Paragraph(html.escape(experience.rewritten_text), styles["Body"]))
        story.extend(
            [
                Paragraph("8. Skills priorizadas", styles["Heading1"]),
                _pdf_table(["Rank", "Skill", "Categoría", "Evidencia"], [[str(s.priority_rank), s.name, str(s.category), str(s.evidence_status)] for s in package.prioritized_skills]),
                Paragraph("9. Keywords ATS", styles["Heading1"]),
                _pdf_table(["Keyword", "Prioridad", "Frecuencia", "Respaldada"], [[k.keyword, str(k.priority), str(k.frequency_in_jobs), _yes_no(package, k.supported_by_candidate)] for k in package.ats_keywords]),
                Paragraph("10. Compatibilidad por vacante", styles["Heading1"]),
                _pdf_table(["Vacante", "Empresa", "Score", "Banda"], _compatibility_rows(package)),
                Paragraph("11. Fortalezas", styles["Heading1"]),
            ]
        )
        story.extend(_pdf_list(package.key_strengths, styles))
        story.append(Paragraph("12. Brechas", styles["Heading1"]))
        story.extend(_pdf_list(package.critical_gaps, styles))
        story.append(Paragraph("13. Recomendaciones", styles["Heading1"]))
        story.extend(_pdf_list(package.strategic_recommendations, styles))
        story.extend(
            [
                Paragraph("14. Auditoría LinkedIn", styles["Heading1"]),
                _pdf_table(["Componente", "Peso", "Score"], _audit_rows(package, "linkedin")),
                Paragraph("15. Auditoría ATS", styles["Heading1"]),
                _pdf_table(["Componente", "Peso", "Score"], _audit_rows(package, "ats")),
                Paragraph("16. Metodología y disclaimer", styles["Heading1"]),
                Paragraph(html.escape(_disclaimer(package)), styles["Body"]),
            ]
        )
        doc.build(story, onFirstPage=_pdf_footer(package), onLaterPages=_pdf_footer(package))
        return output.getvalue()

    def export_zip(self, package: FinalDeliverablePackage, *, banner_image_bytes: bytes | None = None) -> bytes:
        """Export the complete ZIP package."""
        banner_bytes = banner_image_bytes if package.banner_included else None
        files: dict[str, bytes] = {
            "README.txt": self.export_readme(package),
            "manifest.json": self.export_manifest(package, _zip_file_list(package, bool(banner_bytes))),
            "linkedin-profile.md": self.export_markdown(package),
            "linkedin-profile.html": self.export_html(package, banner_image_bytes=banner_bytes, embed_banner=False),
            "linkedin-profile.docx": self.export_docx(package, banner_image_bytes=banner_bytes),
            "linkedin-profile.pdf": self.export_pdf(package, banner_image_bytes=banner_bytes),
            "data/compatibility-summary.json": self.export_compatibility_summary(package),
            "data/audit-summary.json": self.export_audit_summary(package),
        }
        if banner_bytes:
            files[BANNER_ZIP_FILENAME] = banner_bytes
        output = BytesIO()
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for name, data in files.items():
                archive.writestr(f"{ZIP_ROOT}/{name}", data)
        return output.getvalue()

    def export_manifest(self, package: FinalDeliverablePackage, included_files: list[str]) -> bytes:
        """Export manifest JSON bytes."""
        payload = {
            "package_version": package.package_version,
            "export_version": FINAL_PACKAGE_EXPORT_VERSION,
            "created_at": package.generated_at.isoformat(),
            "output_language": package.output_language,
            "content_source": package.content_source,
            "included_files": included_files,
            "banner_included": package.banner_included,
            "linkedin_score": _linkedin_score(package),
            "ats_score": _ats_score(package),
            "average_compatibility_score": package.compatibility_report.average_compatibility_score,
            "methodology_versions": {
                "compatibility": package.compatibility_report.methodology_version,
                "audit": package.audit_report.methodology_version,
            },
            "disclaimer": _disclaimer(package),
        }
        return _json_bytes(payload)

    def export_readme(self, package: FinalDeliverablePackage) -> bytes:
        """Export ZIP README text bytes."""
        if package.output_language == OutputLanguage.EN.value:
            lines = [
                "Astrogato Vector professional package",
                "",
                f"Generated at: {_format_datetime(package.generated_at)}",
                "",
                "Use the Markdown file for copying text, HTML for local review/printing, DOCX for editing, PDF for sharing and ZIP as the complete bundle.",
                "These documents require human review. Scores do not guarantee hiring outcomes.",
                "Missing keywords must not be presented as experience unless supported by evidence.",
                "The LinkedIn banner was generated locally when included.",
            ]
        else:
            lines = [
                "Paquete profesional de Astrogato Vector",
                "",
                f"Fecha de generación: {_format_datetime(package.generated_at)}",
                "",
                "Usa Markdown para copiar texto, HTML para revisión local/impresión, DOCX para edición, PDF para compartir y ZIP como paquete completo.",
                "Los documentos requieren revisión humana. Los scores no garantizan contratación.",
                "Las keywords faltantes no deben presentarse como experiencia si no están respaldadas por evidencia.",
                "El banner de LinkedIn fue generado localmente cuando está incluido.",
            ]
        return "\n".join(lines).encode("utf-8")

    def export_compatibility_summary(self, package: FinalDeliverablePackage) -> bytes:
        """Export reduced compatibility JSON."""
        payload = [
            {
                "job_index": job.job_index,
                "title": job.job_title,
                "company": job.company,
                "score": job.compatibility_score,
                "band": _band_label(package, job.compatibility_band),
                "strengths": job.strengths,
                "critical_gaps": job.critical_gaps,
                "recommendations": job.recommendations,
            }
            for job in package.compatibility_report.job_compatibilities
        ]
        return _json_bytes(payload)

    def export_audit_summary(self, package: FinalDeliverablePackage) -> bytes:
        """Export reduced audit JSON."""
        high_findings = [
            finding.model_dump(mode="json")
            for finding in package.audit_report.findings
            if finding.severity in {"high", "critical"}
        ]
        payload = {
            "linkedin_score": _linkedin_score(package),
            "ats_score": _ats_score(package),
            "overall_score": package.audit_report.overall_score,
            "strengths": package.key_strengths,
            "risks": [finding.title for finding in package.audit_report.risks],
            "quick_wins": package.quick_wins,
            "recommendations": package.strategic_recommendations,
            "high_or_critical_findings": high_findings,
        }
        return _json_bytes(payload)


def _markdown_compatibility(package: FinalDeliverablePackage) -> list[str]:
    lines = [
        "",
        "## 10. Compatibilidad por vacante",
        "",
        "| Vacante | Empresa | Score | Banda | Brechas críticas |",
        "| --- | --- | ---: | --- | --- |",
    ]
    for job in package.compatibility_report.job_compatibilities:
        lines.append(
            "| "
            f"{job.job_index}: {_md_cell(job.job_title)} | {_md_cell(job.company or '')} | "
            f"{job.compatibility_score:.1f} | {_md_cell(_band_label(package, job.compatibility_band))} | "
            f"{_md_cell(_join(job.critical_gaps))} |"
        )
    return lines


def _markdown_lists(package: FinalDeliverablePackage) -> list[str]:
    return [
        "",
        "## 11. Fortalezas",
        "",
        *_markdown_bullets(package.key_strengths),
        "",
        "## 12. Brechas",
        "",
        *_markdown_bullets(package.critical_gaps),
        "",
        "## 13. Recomendaciones",
        "",
        *_markdown_bullets(package.strategic_recommendations),
    ]


def _markdown_audits(package: FinalDeliverablePackage) -> list[str]:
    return [
        "",
        "## 14. Auditoría LinkedIn",
        "",
        f"Score: {_linkedin_score(package):.1f}/100",
        "",
        *_markdown_component_table(package, "linkedin"),
        "",
        "## 15. Auditoría ATS",
        "",
        f"Score: {_ats_score(package):.1f}/100",
        "",
        *_markdown_component_table(package, "ats"),
    ]


def _markdown_methodology(package: FinalDeliverablePackage) -> list[str]:
    return [
        "",
        "## 16. Metodología y disclaimer",
        "",
        _disclaimer(package),
    ]


def _markdown_component_table(package: FinalDeliverablePackage, audit_type: str) -> list[str]:
    lines = ["| Componente | Peso | Score |", "| --- | ---: | ---: |"]
    for name, weight, score in _audit_rows(package, audit_type):
        lines.append(f"| {_md_cell(name)} | {_md_cell(weight)} | {_md_cell(score)} |")
    return lines


def _html_css() -> str:
    return """
body { margin: 0; background: #f6f8fb; color: #1f2937; font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; line-height: 1.55; }
main { max-width: 980px; margin: 0 auto; padding: 32px 20px 56px; }
h1, h2, h3 { color: #14365c; margin: 0 0 12px; }
h1 { font-size: 2rem; }
h2 { font-size: 1.3rem; border-bottom: 2px solid #f28c28; padding-bottom: 6px; margin-top: 28px; }
p, li, td, th { font-size: 0.98rem; }
.cover, section, .score-card { background: #fff; border: 1px solid #d8dee9; border-radius: 8px; padding: 20px; margin: 16px 0; }
.banner { width: 100%; height: auto; border-radius: 6px; border: 1px solid #d8dee9; margin: 8px 0 16px; }
table { width: 100%; border-collapse: collapse; margin: 12px 0; }
th { background: #14365c; color: #fff; text-align: left; }
td, th { border: 1px solid #d8dee9; padding: 8px; vertical-align: top; }
.badge { display: inline-block; border: 1px solid #f28c28; color: #14365c; border-radius: 4px; padding: 2px 6px; margin: 2px; }
dl { display: grid; grid-template-columns: minmax(130px, 220px) 1fr; gap: 8px 16px; }
dt { font-weight: 700; color: #14365c; }
@media (max-width: 680px) { main { padding: 20px 12px; } dl { display: block; } table { display: block; overflow-x: auto; } }
@media print { body { background: #fff; } section, .cover { break-inside: avoid; border: 0; padding: 0; } h2 { break-after: avoid; } }
""".strip()


def _html_section(title: str, content: str) -> str:
    return f"<section><h2>{html.escape(title)}</h2>{content}</section>"


def _html_kv(key: str, value: Any) -> str:
    return f"<dt>{html.escape(str(key))}</dt><dd>{html.escape(str(value))}</dd>"


def _html_list(values: list[str]) -> str:
    items = values or ["No disponible"]
    return "<ul>" + "".join(f"<li>{html.escape(str(item))}</li>" for item in items) + "</ul>"


def _html_experiences(package: FinalDeliverablePackage) -> str:
    return "".join(
        f"<article><h3>{html.escape(experience.employer)} - {html.escape(experience.suggested_role_title)}</h3>"
        f"<p>{html.escape(experience.rewritten_text)}</p></article>"
        for experience in package.experiences
    )


def _html_skills_table(package: FinalDeliverablePackage) -> str:
    rows = [
        [str(skill.priority_rank), skill.name, str(skill.category), str(skill.evidence_status), _join(skill.recommended_placement)]
        for skill in sorted(package.prioritized_skills, key=lambda item: item.priority_rank)
    ]
    return _html_table(["Rank", "Skill", "Categoría", "Evidencia", "Ubicación"], rows)


def _html_keywords_table(package: FinalDeliverablePackage) -> str:
    rows = [
        [
            keyword.keyword,
            str(keyword.priority),
            str(keyword.frequency_in_jobs),
            _yes_no(package, keyword.supported_by_candidate),
            _join(keyword.recommended_sections),
        ]
        for keyword in package.ats_keywords
    ]
    return _html_table(["Keyword", "Prioridad", "Frecuencia", "Respaldada", "Secciones"], rows)


def _html_compatibility_table(package: FinalDeliverablePackage) -> str:
    return _html_table(["Vacante", "Empresa", "Score", "Banda"], _compatibility_rows(package))


def _html_audit_components(package: FinalDeliverablePackage, audit_type: str) -> str:
    return _html_table(["Componente", "Peso", "Score"], _audit_rows(package, audit_type))


def _html_table(headers: list[str], rows: list[list[Any]]) -> str:
    head = "<tr>" + "".join(f"<th>{html.escape(header)}</th>" for header in headers) + "</tr>"
    body = "".join(
        "<tr>" + "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row) + "</tr>"
        for row in rows
    )
    return f"<table>{head}{body}</table>"


def _setup_docx(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2"):
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = RGBColor(20, 54, 92)


def _docx_heading(document: Document, text: str, *, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _docx_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def _docx_bullets(document: Document, values: list[str]) -> None:
    for value in values or ["No disponible"]:
        document.add_paragraph(str(value), style="List Bullet")


def _docx_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    document.add_paragraph()


def _docx_add_banner(document: Document, banner_image_bytes: bytes) -> None:
    image = BytesIO(banner_image_bytes)
    try:
        document.add_picture(image, width=Inches(6.4))
    except Exception:
        _docx_paragraph(document, "Banner PNG no insertado en DOCX por un error de lectura de imagen.")


def _docx_audit(document: Document, package: FinalDeliverablePackage, audit_type: str) -> None:
    _docx_table(document, ["Componente", "Peso", "Score"], _audit_rows(package, audit_type))


def _scrub_docx_core_properties(document: Document) -> None:
    props = document.core_properties
    props.author = "Astrogato Vector"
    props.last_modified_by = "Astrogato Vector"
    props.comments = "Generated with Astrogato Vector"
    props.subject = "Professional positioning package"
    props.title = "Astrogato Vector professional package"
    props.keywords = ""


def _pdf_styles() -> dict[str, ParagraphStyle]:
    font_name = _register_pdf_font()
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "AVTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#14365c"),
            alignment=TA_LEFT,
            spaceAfter=16,
        ),
        "Subtitle": ParagraphStyle("AVSubtitle", parent=base["BodyText"], fontName=font_name, fontSize=12, leading=16, spaceAfter=12),
        "Heading1": ParagraphStyle("AVHeading1", parent=base["Heading1"], fontName=font_name, fontSize=14, leading=18, textColor=colors.HexColor("#14365c"), spaceBefore=16, spaceAfter=8),
        "Heading2": ParagraphStyle("AVHeading2", parent=base["Heading2"], fontName=font_name, fontSize=12, leading=15, textColor=colors.HexColor("#14365c"), spaceBefore=10, spaceAfter=6),
        "Body": ParagraphStyle("AVBody", parent=base["BodyText"], fontName=font_name, fontSize=9.5, leading=13, spaceAfter=7),
        "Small": ParagraphStyle("AVSmall", parent=base["BodyText"], fontName=font_name, fontSize=8, leading=10, textColor=colors.HexColor("#5b6472")),
    }


def _register_pdf_font() -> str:
    path = find_font_path(DEFAULT_FONT_NAMES, bold=False)
    if path:
        try:
            pdfmetrics.registerFont(TTFont("AstrogatoSans", path))
            return "AstrogatoSans"
        except Exception:
            pass
    return "Helvetica"


def _pdf_table(headers: list[str], rows: list[list[Any]]) -> Table:
    font_name = _register_pdf_font()
    cell_style = ParagraphStyle("AVTableCell", fontName=font_name, fontSize=7.4, leading=9.2)
    header_style = ParagraphStyle("AVTableHeader", fontName=font_name, fontSize=7.4, leading=9.2, textColor=colors.white)
    data = [
        [Paragraph(html.escape(str(header)), header_style) for header in headers],
        *[[Paragraph(html.escape(str(cell)), cell_style) for cell in row] for row in rows],
    ]
    usable_width = LETTER[0] - (1.44 * inch)
    table = Table(data, colWidths=[usable_width / len(headers)] * len(headers), repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#14365c")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d8dee9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 7.6),
                ("LEADING", (0, 0), (-1, -1), 9.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _pdf_list(values: list[str], styles: dict[str, ParagraphStyle]) -> list[Any]:
    return [Paragraph(html.escape(f"- {value}"), styles["Body"]) for value in (values or ["No disponible"])]


def _pdf_banner(banner_image_bytes: bytes) -> Image:
    image_buffer = BytesIO(banner_image_bytes)
    with PILImage.open(BytesIO(banner_image_bytes)) as image:
        width, height = image.size
    target_width = 6.4 * inch
    target_height = target_width * height / width
    return Image(image_buffer, width=target_width, height=target_height)


def _pdf_footer(package: FinalDeliverablePackage):
    def draw_footer(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#5b6472"))
        canvas.drawString(0.72 * inch, 0.42 * inch, "Generated with Astrogato Vector")
        canvas.drawRightString(7.78 * inch, 0.42 * inch, f"{doc.page}")
        canvas.restoreState()

    return draw_footer


def _compatibility_rows(package: FinalDeliverablePackage) -> list[list[str]]:
    return [
        [
            f"{job.job_index}: {job.job_title}",
            job.company or "",
            f"{job.compatibility_score:.1f}",
            _band_label(package, job.compatibility_band),
        ]
        for job in package.compatibility_report.job_compatibilities
    ]


def _audit_rows(package: FinalDeliverablePackage, audit_type: str) -> list[list[str]]:
    audit = package.audit_report.linkedin_positioning if audit_type == "linkedin" else package.audit_report.ats_estimation
    if audit is None:
        return []
    return [[component.name, f"{component.weight:.0%}", f"{component.score:.1f}"] for component in audit.components]


def _zip_file_list(package: FinalDeliverablePackage, include_banner: bool) -> list[str]:
    files = [
        "README.txt",
        "manifest.json",
        "linkedin-profile.md",
        "linkedin-profile.html",
        "linkedin-profile.docx",
        "linkedin-profile.pdf",
        "data/compatibility-summary.json",
        "data/audit-summary.json",
    ]
    if include_banner:
        files.append(BANNER_ZIP_FILENAME)
    return files


def _band_label(package: FinalDeliverablePackage, value: object) -> str:
    key = str(getattr(value, "value", value))
    labels = COMPATIBILITY_BAND_LABELS_EN if package.output_language == OutputLanguage.EN.value else COMPATIBILITY_BAND_LABELS_ES
    return labels.get(key, key)


def _linkedin_score(package: FinalDeliverablePackage) -> float | None:
    return package.audit_report.linkedin_positioning.score if package.audit_report.linkedin_positioning else None


def _ats_score(package: FinalDeliverablePackage) -> float | None:
    return package.audit_report.ats_estimation.score if package.audit_report.ats_estimation else None


def _disclaimer(package: FinalDeliverablePackage) -> str:
    language = str(package.output_language)
    return FINAL_PACKAGE_DISCLAIMERS.get(language, FINAL_AUDIT_DISCLAIMER)


def _label(package: FinalDeliverablePackage, key: str) -> str:
    en = package.output_language == OutputLanguage.EN.value
    labels = {
        "date": ("Date", "Fecha"),
        "language": ("Language", "Idioma"),
        "identity": ("Professional identity", "Identidad profesional"),
        "roles": ("Target roles", "Roles objetivo"),
        "banner_primary": ("Primary line", "Línea principal"),
        "banner_specialty": ("Specialty line", "Especialidades"),
        "banner_supporting": ("Supporting line", "Línea de apoyo"),
        "not_available": ("Not available", "No disponible"),
        "yes": ("Yes", "Sí"),
        "no": ("No", "No"),
    }
    return labels[key][0 if en else 1]


def _yes_no(package: FinalDeliverablePackage, value: bool) -> str:
    return _label(package, "yes") if value else _label(package, "no")


def _format_datetime(value: datetime) -> str:
    return value.isoformat(timespec="seconds")


def _metadata_line(label: str, value: Any) -> str:
    return f"**{label}:** {value}"


def _join(values: list[str]) -> str:
    cleaned = [str(value).strip() for value in values if str(value).strip()]
    return ", ".join(cleaned) if cleaned else "No disponible"


def _markdown_bullets(values: list[str]) -> list[str]:
    return [f"- {value}" for value in values] if values else ["- No disponible"]


def _md_cell(value: Any) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def _json_bytes(payload: Any) -> bytes:
    return json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True).encode("utf-8")
