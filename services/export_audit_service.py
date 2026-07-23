"""Audit generated export bytes before exposing downloads."""

from __future__ import annotations

import io
import zipfile

from docx import Document
from pypdf import PdfReader

from schemas.deliverable_models import (
    FINAL_PACKAGE_REQUIRED_SECTIONS,
    MAX_FINAL_ZIP_SIZE_MB,
    MAX_INDIVIDUAL_EXPORT_SIZE_MB,
    ExportAuditResult,
)

_FORBIDDEN_ZIP_SUFFIXES = (".env", ".log", ".tmp")
_FORBIDDEN_ZIP_NAMES = {"cv.pdf", "cv.docx", "resume.pdf", "resume.docx", "prompt.txt", "openai-response.json"}
_TARGETED_CV_FORBIDDEN_TERMS = (
    "compatibility_score",
    "request_id",
    "input_tokens",
    "output_tokens",
    "prompt",
    "raw_response",
    "openai",
    "score",
    "brecha",
    "auditoría",
    "auditoria",
    "evidencia interna",
)
_TARGETED_CV_REQUIRED_SECTIONS = ("Perfil profesional", "Competencias clave", "Experiencia profesional")


class ExportAuditService:
    """Validate export bytes by format."""

    def audit_markdown(self, data: bytes) -> ExportAuditResult:
        findings: list[str] = []
        warnings: list[str] = []
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "markdown", findings)
        text = _decode_utf8(data, findings, "markdown")
        if text:
            for section in FINAL_PACKAGE_REQUIRED_SECTIONS:
                if section not in text:
                    findings.append(f"markdown: falta sección {section}.")
            if "<html" in text.casefold():
                findings.append("markdown: no debe contener HTML incrustado.")
        return ExportAuditResult(passed=not findings, findings=findings, warnings=warnings)

    def audit_html(self, data: bytes) -> ExportAuditResult:
        findings: list[str] = []
        text = _decode_utf8(data, findings, "html")
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "html", findings)
        if text:
            lowered = text.casefold()
            if "<!doctype html>" not in lowered:
                findings.append("html: falta doctype HTML5.")
            if "<html" not in lowered:
                findings.append("html: falta etiqueta html.")
            if "charset=\"utf-8\"" not in lowered and "charset='utf-8'" not in lowered:
                findings.append("html: falta charset UTF-8.")
            if "<script" in lowered:
                findings.append("html: no debe contener scripts.")
            if "http://" in lowered or "https://" in lowered:
                findings.append("html: no debe contener recursos remotos.")
            if "</html>" not in lowered or "</body>" not in lowered:
                findings.append("html: estructura no cerrada.")
            for section in FINAL_PACKAGE_REQUIRED_SECTIONS:
                if section.casefold() not in lowered:
                    findings.append(f"html: falta sección {section}.")
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_docx(self, data: bytes) -> ExportAuditResult:
        findings: list[str] = []
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "docx", findings)
        if not zipfile.is_zipfile(io.BytesIO(data)):
            findings.append("docx: firma ZIP inválida.")
            return ExportAuditResult(passed=False, findings=findings)
        try:
            document = Document(io.BytesIO(data))
        except Exception as exc:
            findings.append(f"docx: no se pudo abrir con python-docx: {exc}")
            return ExportAuditResult(passed=False, findings=findings)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if not text.strip():
            findings.append("docx: no contiene párrafos.")
        for section in FINAL_PACKAGE_REQUIRED_SECTIONS:
            if section not in text:
                findings.append(f"docx: falta sección {section}.")
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_pdf(self, data: bytes) -> ExportAuditResult:
        findings: list[str] = []
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "pdf", findings)
        if not data.startswith(b"%PDF"):
            findings.append("pdf: firma inválida.")
            return ExportAuditResult(passed=False, findings=findings)
        if len(data) < 1024:
            findings.append("pdf: tamaño demasiado pequeño.")
        try:
            reader = PdfReader(io.BytesIO(data))
            if len(reader.pages) < 1:
                findings.append("pdf: no contiene páginas.")
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if len(text.strip()) < 100:
                findings.append("pdf: no contiene texto seleccionable suficiente.")
            for section in FINAL_PACKAGE_REQUIRED_SECTIONS:
                if section not in text:
                    findings.append(f"pdf: falta sección {section}.")
        except Exception as exc:
            findings.append(f"pdf: no se pudo abrir: {exc}")
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_zip(self, data: bytes) -> ExportAuditResult:
        findings: list[str] = []
        _validate_size(data, MAX_FINAL_ZIP_SIZE_MB, "zip", findings)
        try:
            with zipfile.ZipFile(io.BytesIO(data), "r") as archive:
                names = archive.namelist()
                lowered = {name.casefold() for name in names}
                required = {
                    "astrogato-vector-paquete-profesional/readme.txt",
                    "astrogato-vector-paquete-profesional/manifest.json",
                    "astrogato-vector-paquete-profesional/linkedin-profile.md",
                    "astrogato-vector-paquete-profesional/linkedin-profile.html",
                    "astrogato-vector-paquete-profesional/linkedin-profile.docx",
                    "astrogato-vector-paquete-profesional/linkedin-profile.pdf",
                    "astrogato-vector-paquete-profesional/data/compatibility-summary.json",
                    "astrogato-vector-paquete-profesional/data/audit-summary.json",
                }
                for required_name in required:
                    if required_name not in lowered:
                        findings.append(f"zip: falta {required_name}.")
                for name in names:
                    normalized = name.replace("\\", "/")
                    if normalized.startswith("/") or ":" in normalized or ".." in normalized.split("/"):
                        findings.append(f"zip: ruta insegura {name}.")
                    basename = normalized.rsplit("/", 1)[-1].casefold()
                    if basename in _FORBIDDEN_ZIP_NAMES or basename.endswith(_FORBIDDEN_ZIP_SUFFIXES):
                        findings.append(f"zip: archivo prohibido {name}.")
        except Exception as exc:
            findings.append(f"zip: no se pudo abrir: {exc}")
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_all(self, exports: dict[str, bytes]) -> ExportAuditResult:
        """Audit all expected exports and merge findings."""
        findings: list[str] = []
        warnings: list[str] = []
        checks = {
            "markdown": self.audit_markdown,
            "html": self.audit_html,
            "docx": self.audit_docx,
            "pdf": self.audit_pdf,
            "zip": self.audit_zip,
        }
        for key, audit_fn in checks.items():
            result = audit_fn(exports.get(key, b""))
            findings.extend(result.findings)
            warnings.extend(result.warnings)
        return ExportAuditResult(passed=not findings, findings=findings, warnings=warnings)

    def audit_targeted_cv_markdown(self, data: bytes) -> ExportAuditResult:
        """Audit one targeted CV Markdown export."""
        findings: list[str] = []
        warnings: list[str] = []
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "targeted_cv_markdown", findings)
        text = _decode_utf8(data, findings, "targeted_cv_markdown")
        if text:
            if "<html" in text.casefold():
                findings.append("targeted_cv_markdown: no debe contener HTML.")
            for section in _TARGETED_CV_REQUIRED_SECTIONS:
                if section not in text:
                    findings.append(f"targeted_cv_markdown: falta sección {section}.")
            _validate_no_targeted_cv_internal_terms(text, "targeted_cv_markdown", findings)
        return ExportAuditResult(passed=not findings, findings=findings, warnings=warnings)

    def audit_targeted_cv_docx(self, data: bytes) -> ExportAuditResult:
        """Audit one targeted CV DOCX export."""
        findings: list[str] = []
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "targeted_cv_docx", findings)
        if not zipfile.is_zipfile(io.BytesIO(data)):
            findings.append("targeted_cv_docx: firma ZIP inválida.")
            return ExportAuditResult(passed=False, findings=findings)
        try:
            document = Document(io.BytesIO(data))
        except Exception as exc:
            findings.append(f"targeted_cv_docx: no se pudo abrir con python-docx: {exc}")
            return ExportAuditResult(passed=False, findings=findings)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if not text.strip():
            findings.append("targeted_cv_docx: no contiene párrafos.")
        for section in _TARGETED_CV_REQUIRED_SECTIONS:
            if section.upper() not in text and section not in text:
                findings.append(f"targeted_cv_docx: falta sección {section}.")
        if document.tables:
            findings.append("targeted_cv_docx: no debe usar tablas para el layout del CV.")
        if document.inline_shapes:
            findings.append("targeted_cv_docx: no debe incluir imágenes.")
        _validate_no_targeted_cv_internal_terms(text, "targeted_cv_docx", findings)
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_targeted_cv_pdf(self, data: bytes) -> ExportAuditResult:
        """Audit one targeted CV PDF export."""
        findings: list[str] = []
        _validate_size(data, MAX_INDIVIDUAL_EXPORT_SIZE_MB, "targeted_cv_pdf", findings)
        if not data.startswith(b"%PDF"):
            findings.append("targeted_cv_pdf: firma inválida.")
            return ExportAuditResult(passed=False, findings=findings)
        try:
            reader = PdfReader(io.BytesIO(data))
            if len(reader.pages) < 1:
                findings.append("targeted_cv_pdf: no contiene páginas.")
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if len(text.strip()) < 100:
                findings.append("targeted_cv_pdf: no contiene texto seleccionable suficiente.")
            for section in _TARGETED_CV_REQUIRED_SECTIONS:
                if section.upper() not in text and section not in text:
                    findings.append(f"targeted_cv_pdf: falta sección {section}.")
            _validate_no_targeted_cv_internal_terms(text, "targeted_cv_pdf", findings)
        except Exception as exc:
            findings.append(f"targeted_cv_pdf: no se pudo abrir: {exc}")
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_targeted_cv_zip(self, data: bytes) -> ExportAuditResult:
        """Audit the targeted CV ZIP package."""
        findings: list[str] = []
        _validate_size(data, MAX_FINAL_ZIP_SIZE_MB, "targeted_cv_zip", findings)
        try:
            with zipfile.ZipFile(io.BytesIO(data), "r") as archive:
                names = archive.namelist()
                lowered = {name.casefold() for name in names}
                if "targeted-cvs/readme.txt" not in lowered:
                    findings.append("targeted_cv_zip: falta targeted-cvs/README.txt.")
                if "targeted-cvs/manifest.json" not in lowered:
                    findings.append("targeted_cv_zip: falta targeted-cvs/manifest.json.")
                vacancy_folders = sorted(
                    {
                        name.split("/")[1]
                        for name in names
                        if name.startswith("targeted-cvs/vacancy-") and len(name.split("/")) >= 3
                    }
                )
                if not vacancy_folders:
                    findings.append("targeted_cv_zip: no contiene carpetas vacancy-XX.")
                for folder in vacancy_folders:
                    for basename in ("cv.md", "cv.docx", "cv.pdf", "review-summary.json"):
                        required = f"targeted-cvs/{folder}/{basename}".casefold()
                        if required not in lowered:
                            findings.append(f"targeted_cv_zip: falta {required}.")
                for name in names:
                    normalized = name.replace("\\", "/")
                    if normalized.startswith("/") or ":" in normalized or ".." in normalized.split("/"):
                        findings.append(f"targeted_cv_zip: ruta insegura {name}.")
                    if not normalized.startswith("targeted-cvs/"):
                        findings.append(f"targeted_cv_zip: archivo fuera de raíz {name}.")
                    basename = normalized.rsplit("/", 1)[-1].casefold()
                    if basename in {"prompt.txt", "openai-response.json", "raw-cv.txt", "original-cv.pdf"}:
                        findings.append(f"targeted_cv_zip: archivo prohibido {name}.")
                for name in names:
                    lowered_name = name.casefold()
                    payload = archive.read(name)
                    if lowered_name.endswith("/cv.md"):
                        findings.extend(self.audit_targeted_cv_markdown(payload).findings)
                    elif lowered_name.endswith("/cv.docx"):
                        findings.extend(self.audit_targeted_cv_docx(payload).findings)
                    elif lowered_name.endswith("/cv.pdf"):
                        findings.extend(self.audit_targeted_cv_pdf(payload).findings)
        except Exception as exc:
            findings.append(f"targeted_cv_zip: no se pudo abrir: {exc}")
        return ExportAuditResult(passed=not findings, findings=findings)

    def audit_targeted_cv_all(self, exports: dict[str, bytes]) -> ExportAuditResult:
        """Audit all targeted CV export formats."""
        findings: list[str] = []
        warnings: list[str] = []
        checks = {
            "markdown": self.audit_targeted_cv_markdown,
            "docx": self.audit_targeted_cv_docx,
            "pdf": self.audit_targeted_cv_pdf,
            "zip": self.audit_targeted_cv_zip,
        }
        for key, audit_fn in checks.items():
            result = audit_fn(exports.get(key, b""))
            findings.extend(result.findings)
            warnings.extend(result.warnings)
        return ExportAuditResult(passed=not findings, findings=findings, warnings=warnings)


def _decode_utf8(data: bytes, findings: list[str], label: str) -> str:
    if not data:
        findings.append(f"{label}: bytes vacíos.")
        return ""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        findings.append(f"{label}: UTF-8 inválido.")
        return ""


def _validate_size(data: bytes, max_mb: int, label: str, findings: list[str]) -> None:
    if len(data) > max_mb * 1024 * 1024:
        findings.append(f"{label}: supera el tamaño máximo de {max_mb} MB.")


def _validate_no_targeted_cv_internal_terms(text: str, label: str, findings: list[str]) -> None:
    lowered = text.casefold()
    for term in _TARGETED_CV_FORBIDDEN_TERMS:
        if term.casefold() in lowered:
            findings.append(f"{label}: contiene término interno no permitido: {term}.")
